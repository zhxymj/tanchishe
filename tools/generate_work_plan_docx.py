from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "贪吃蛇项目五人详细工作步骤安排.docx"


ACCENT = (46, 116, 181)
ACCENT_DARK = (31, 77, 120)
HEADER_FILL = "E8EEF5"


def set_font(run, size=10.5, bold=False, color=(17, 24, 39), font="微软雅黑"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("贪吃蛇项目五人详细工作步骤安排")
    set_font(run, size=22, bold=True, color=(31, 41, 55))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("C11 + raylib 图形化单机游戏 · 小组协作版")
    set_font(run, size=12, color=(75, 85, 99))


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    run = p.add_run(text)
    set_font(run, size=16 if level == 1 else 13, bold=True, color=ACCENT if level == 1 else ACCENT_DARK)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.65)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("• " + text)
    set_font(run)
    return p


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F3F4F6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    for index, line in enumerate(text.strip().splitlines()):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_font(run, size=9.2, font="Consolas")
    doc.add_paragraph()


def add_table(doc, rows, widths):
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for row_index, row_data in enumerate(rows):
        cells = table.rows[0].cells if row_index == 0 else table.add_row().cells
        for col_index, text in enumerate(row_data):
            cell = cells[col_index]
            cell.width = Cm(widths[col_index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            if row_index == 0:
                shade_cell(cell, HEADER_FILL)
                set_font(run, size=9.5, bold=True, color=(31, 77, 120))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                set_font(run, size=9.2)
    doc.add_paragraph()
    return table


def add_member(doc, title, owner, functions, steps, checks):
    add_heading(doc, title, 2)
    add_table(
        doc,
        [
            ["项目", "内容"],
            ["责任范围", owner],
            ["主要函数", functions],
        ],
        [3.0, 12.0],
    )
    add_body(doc, "详细步骤：")
    for index, item in enumerate(steps, 1):
        add_body(doc, f"{index}. {item}")
    add_body(doc, "自测重点：")
    for item in checks:
        add_bullet(doc, item)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    add_title(doc)
    add_heading(doc, "一、项目当前状态", 1)
    add_body(doc, "本项目已经从控制台字符版调整为 C11 + raylib 图形化单机贪吃蛇。当前代码拆分为 main、game、snake、food、ui 五个模块，配套 build.bat、CMakeLists.txt 和 README.md。")
    add_table(
        doc,
        [
            ["检查项", "当前状态"],
            ["实现语言", "C11"],
            ["图形库", "raylib 6.0，安装路径 D:\\raylib"],
            ["游戏类型", "电脑键盘控制的窗口化单机游戏"],
            ["编译方式", "build.bat、gcc 命令、CMake（英文路径更稳定）"],
            ["仓库地址", "https://github.com/zhxymj/tanchishe"],
        ],
        [3.2, 11.8],
    )

    add_heading(doc, "二、统一开发流程", 1)
    add_body(doc, "每次写代码前先同步，写完后先编译，再提交。当前代码已经拆分为 main、game、snake、food、ui 五个模块，成员之间按文件协作。")
    add_code(
        doc,
        r"""
git pull origin main
.\build.bat
git diff
git add src\自己负责的文件.c src\自己负责的文件.h
git commit -m "说明本次完成的功能"
git push origin main
""",
    )

    members = [
        (
            "三、成员1：主程序框架与状态管理",
            "窗口初始化、游戏初始化、状态切换、暂停继续、重开和固定时间步进。",
            "src/main.c, src/game.c, src/game.h",
            [
                "确认窗口大小、帧率和音频设备初始化正确。",
                "维护开始前、运行中、暂停、游戏结束四种状态。",
                "保证蛇移动使用固定时间步进，不受帧率波动影响。",
                "和其他成员联调按钮、键盘和游戏结束流程。",
            ],
            ["Space/P 可以暂停继续。", "R 可以重新开始。", "游戏结束后 Enter 可以重开。"],
        ),
        (
            "四、成员2：蛇身移动与碰撞规则",
            "蛇身数据、方向控制、禁止反向、移动增长、撞墙和撞自身。",
            "src/snake.c, src/snake.h",
            [
                "维护 body[0] 为蛇头、后续元素为蛇身的约定。",
                "处理方向键和 WASD 传入后的方向队列。",
                "吃到食物时增长，没吃到时正常移动。",
                "判断撞墙和撞自身，并切换到游戏结束状态。",
            ],
            ["不能直接反向移动。", "吃食物后长度增加。", "撞墙和撞自身会结束游戏。"],
        ),
        (
            "五、成员3：食物、计分、难度和最高分",
            "食物刷新、得分、速度等级、最高分文件和食物视觉反馈。",
            "src/food.c, src/food.h",
            [
                "保证食物随机生成在棋盘内部。",
                "避免食物生成在蛇身上。",
                "吃到食物后加分、触发动画和音效反馈。",
                "根据分数提升等级并逐步加快速度。",
                "保存和读取 snake_raylib_highscore.dat。",
            ],
            ["食物不会和蛇重合。", "分数和最高分显示正确。", "速度条随等级变化。"],
        ),
        (
            "六、成员4：棋盘、蛇身和背景视觉",
            "现代街机风格背景、棋盘边界、网格、蛇身层次和蛇头方向表现。",
            "src/ui.c, src/ui.h 中的绘制函数",
            [
                "保持 960x720 下棋盘和面板不重叠。",
                "绘制深色背景、轻微网格和棋盘阴影。",
                "绘制圆角蛇身、颜色层次和高光。",
                "根据方向绘制蛇头眼睛，让朝向清楚。",
            ],
            ["棋盘居中且边界清晰。", "蛇身不是字符块。", "蛇头方向可识别。"],
        ),
        (
            "七、成员5：面板、按钮、弹窗和构建说明",
            "鼠标按钮、右侧信息面板、暂停/结束弹窗、文字适配和构建脚本。",
            "build.bat, CMakeLists.txt, README.md, TASKS.md, CONTRIBUTING.md，并配合维护 src/ui.c 的按钮区域",
            [
                "维护 Start、Pause、Restart 按钮的 hover 和 pressed 状态。",
                "绘制分数、最高分、等级、速度和状态面板。",
                "暂停和游戏结束时绘制半透明覆盖层弹窗。",
                "保证 build.bat 默认使用 D:\\raylib。",
                "在 README 中维护正确编译和运行说明。",
            ],
            ["按钮可鼠标点击。", "弹窗文字和按钮不重叠。", "README 命令能编译运行。"],
        ),
    ]

    for member in members:
        add_member(doc, *member)

    add_heading(doc, "八、提交前最终检查清单", 1)
    for item in [
        "源码能通过 gcc 编译。",
        "不要提交 snake_raylib.exe、build/ 或最高分数据文件。",
        "README、TASKS、CONTRIBUTING 与当前 raylib 版本一致。",
        "每个成员都能在报告中说明自己负责的函数和代码逻辑。",
        "演示时能完成开始、移动、吃食物、暂停、重开和游戏结束流程。",
    ]:
        add_bullet(doc, item)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
